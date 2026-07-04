from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
LOGO = ROOT / "media" / "logo-upt.png"
TODAY = date(2026, 7, 4).strftime("%d/%m/%Y")
PROJECT = "Drop Nexus - Replicador de Datos entre Bases de Datos"
SYSTEM = "Drop Nexus"
COURSE = "SI783 - Proyecto de Sistemas"
DOCENTE = "Docente del curso"
TEAM = "Equipo del proyecto"
YEAR = "2026"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)


def styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 15.5, "1F4E79", 14, 8),
        ("Heading 2", 12.5, "2E74B5", 10, 5),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def cover(doc: Document, doc_title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.0))
    for line, size in [
        ("UNIVERSIDAD PRIVADA DE TACNA", 13),
        ("FACULTAD DE INGENIERIA", 12),
        ("Escuela Profesional de Ingenieria de Sistemas", 11),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(size)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Proyecto {PROJECT}")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_title)
    r.bold = True
    r.font.size = Pt(14)
    for label, value in [
        ("Curso", COURSE),
        ("Docente", DOCENTE),
        ("Integrantes", TEAM),
        ("Tacna - Peru", YEAR),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: " if label not in {"Tacna - Peru"} else f"{label} ")
        run.bold = True
        p.add_run(value)
    doc.add_page_break()


def version_and_index(doc: Document, doc_title: str, index: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Sistema {SYSTEM}")
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_title)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Version 1.0").bold = True
    doc.add_heading("CONTROL DE VERSIONES", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Version", "Hecha por", "Revisada por", "Aprobada por", "Fecha", "Motivo"]
    for idx, text in enumerate(headers):
        c = table.rows[0].cells[idx]
        c.text = text
        shade(c, "E8EEF5")
        c.paragraphs[0].runs[0].bold = True
    row = table.add_row().cells
    for idx, text in enumerate(["1.0", TEAM, "Pendiente", "Pendiente", TODAY, "Adaptacion al sistema implementado"]):
        row[idx].text = text
    table_geometry(table, [0.75, 1.25, 1.25, 1.25, 0.9, 2.05])
    doc.add_page_break()
    doc.add_heading("INDICE GENERAL", level=1)
    for item in index:
        doc.add_paragraph(item)
    doc.add_page_break()


def footer(doc: Document, title: str) -> None:
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = f"{SYSTEM} | {title}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(89, 89, 89)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, h in enumerate(headers):
        c = table.rows[0].cells[idx]
        c.text = h
        shade(c, "E8EEF5")
        c.paragraphs[0].runs[0].bold = True
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value
    table_geometry(table, widths)


def build(path: str, title: str, index: list[str], body) -> None:
    doc = Document()
    styles(doc)
    cover(doc, title)
    version_and_index(doc, title, index)
    body(doc)
    footer(doc, title)
    doc.save(ROOT / path)


def factibilidad(doc: Document) -> None:
    doc.add_heading("Informe de Factibilidad", level=1)
    doc.add_heading("1. Descripcion del Proyecto", level=1)
    doc.add_heading("1.1 Nombre del proyecto", level=2)
    para(doc, PROJECT)
    doc.add_heading("1.2 Duracion del proyecto", level=2)
    para(doc, "Ocho semanas academicas para ajuste documental, pruebas, despliegue local y preparacion de exposicion.")
    doc.add_heading("1.3 Descripcion", level=2)
    para(doc, "Drop Nexus es una aplicacion web que permite importar, configurar, replicar y descargar datos entre bases de datos heterogeneas. La app soporta motores relacionales, MongoDB y archivos SQLite, SQL, Excel/CSV y respaldos SQL Server .bak con restauracion configurada.")
    doc.add_heading("1.4 Objetivos", level=2)
    doc.add_heading("1.4.1 Objetivo general", level=3)
    para(doc, "Desarrollar un replicador web que reduzca el trabajo manual de migracion y permita entregar una base modificada descargable.")
    doc.add_heading("1.4.2 Objetivos Especificos", level=3)
    bullets(doc, [
        "Importar archivos SQLite, SQL, MongoDB JSON/NDJSON, Excel/CSV y SQL Server .bak configurado.",
        "Configurar origen y destino, seleccionar tablas, mapear columnas y ejecutar replicas.",
        "Registrar historial, progreso, errores y reintentos.",
        "Descargar la base modificada en JSON, Excel .xlsx o SQLite.",
    ])
    doc.add_heading("2. Riesgos", level=1)
    matrix(doc, ["Riesgo", "Impacto", "Mitigacion"], [
        ["Tipos incompatibles", "Errores de insercion o perdida de precision", "Mapeo por motor y validacion previa."],
        ["Archivo invalido", "Fallo de importacion", "Validacion de extension, tamano, cabecera SQLite y contenido binario."],
        [".bak sin SQL Server", "No se restaura el respaldo", "Variables SQLSERVER_RESTORE_* obligatorias y error claro."],
        ["Credenciales expuestas", "Riesgo de seguridad", "Cifrado AES-256-GCM y exclusion de .env."],
    ], [2.0, 2.0, 2.5])
    doc.add_heading("3. Analisis de la Situacion actual", level=1)
    doc.add_heading("3.1 Planteamiento del problema", level=2)
    para(doc, "La migracion entre motores suele realizarse con scripts aislados y exportaciones manuales. Esto dificulta validar errores, mapear columnas y entregar un resultado verificable.")
    doc.add_heading("3.2 Consideraciones de hardware y software", level=2)
    bullets(doc, ["Node.js 20, React, Fastify y TypeScript.", "Supabase/PostgreSQL para metadatos y catalogos.", "Render como plataforma de despliegue.", "SQL Server externo solo para restauracion .bak."])
    doc.add_heading("4. Estudio de Factibilidad", level=1)
    doc.add_heading("4.1 Factibilidad Tecnica", level=2)
    para(doc, "Es tecnicamente viable porque la app ya cuenta con frontend, backend, adaptadores por motor, catalogos de archivo y render.yaml.")
    doc.add_heading("4.2 Factibilidad economica", level=2)
    matrix(doc, ["Concepto", "Costo estimado"], [["Desarrollo academico", "S/ 0"], ["GitHub y herramientas", "S/ 0"], ["Render/Supabase", "Plan gratuito o bajo costo"], ["SQL Server temporal para .bak", "Variable"]], [3.2, 3.3])
    doc.add_heading("4.2.1 Costos Generales", level=3)
    para(doc, "Uso de equipo personal, internet y herramientas gratuitas.")
    doc.add_heading("4.2.2 Costos operativos durante el desarrollo", level=3)
    para(doc, "No se requiere oficina ni infraestructura fisica dedicada.")
    doc.add_heading("4.2.3 Costos del ambiente", level=3)
    para(doc, "Ambiente local y despliegue en Render con Supabase.")
    doc.add_heading("4.2.4 Costos de personal", level=3)
    para(doc, "Equipo academico del curso.")
    doc.add_heading("4.2.5 Costos totales del desarrollo del sistema", level=3)
    para(doc, "El costo principal es tiempo de desarrollo; los servicios pueden mantenerse en capa gratuita.")
    doc.add_heading("4.3 Factibilidad Operativa", level=2)
    para(doc, "La aplicacion se opera desde navegador y guia la carga, configuracion, replica y descarga.")
    doc.add_heading("4.4 Factibilidad Legal", level=2)
    para(doc, "El sistema debe usarse con datos autorizados; cifra credenciales y aísla catalogos por usuario.")
    doc.add_heading("4.5 Factibilidad Social", level=2)
    para(doc, "Facilita el aprendizaje y reduce errores humanos en migraciones.")
    doc.add_heading("4.6 Factibilidad Ambiental", level=2)
    para(doc, "No requiere infraestructura fisica adicional para el proyecto.")
    doc.add_heading("5. Analisis Financiero", level=1)
    doc.add_heading("5.1 Justificacion de la Inversion", level=2)
    para(doc, "El beneficio supera el costo por reutilizar tecnologias libres y reducir trabajo manual.")
    doc.add_heading("5.1.1 Beneficios del Proyecto", level=3)
    bullets(doc, ["Menos errores manuales.", "Trazabilidad de replicas.", "Descarga de resultados.", "Demostracion academica clara."])
    doc.add_heading("5.1.2 Criterios de Inversion", level=3)
    doc.add_heading("5.1.2.1 Relacion Beneficio/Costo (B/C)", level=3)
    para(doc, "B/C favorable por bajo costo y alto valor educativo.")
    doc.add_heading("5.1.2.2 Valor Actual Neto (VAN)", level=3)
    para(doc, "VAN academico positivo por uso de servicios gratuitos y reutilizacion futura.")
    doc.add_heading("5.1.2.3 Tasa Interna de Retorno (TIR)", level=3)
    para(doc, "No aplica comercialmente; la rentabilidad se expresa en aprendizaje y funcionalidad entregada.")
    doc.add_heading("6. Conclusiones", level=1)
    para(doc, "El proyecto es viable y cumple con la finalidad de replicar datos entre bases y descargar la base modificada.")


def vision(doc: Document) -> None:
    doc.add_heading("Informe de Vision", level=1)
    for title, text in [
        ("1. Introduccion", "Define la vision funcional de Drop Nexus como replicador web de datos entre motores heterogeneos."),
        ("1.1 Proposito", "Alinear el alcance del producto con la app implementada en la carpeta web/."),
        ("1.2 Alcance", "Incluye autenticacion, importacion, configuracion, replicacion, historial, descarga y asistencia integrada."),
        ("1.3 Definiciones, Siglas y Abreviaturas", "Catalogo de archivo: representacion normalizada de una base importada. Replica: proceso de copia y transformacion de filas. .bak: respaldo SQL Server restaurable con infraestructura externa."),
        ("1.4 Referencias", "Codigo fuente web/, README, render.yaml, skill del asistente y extension VS Code."),
        ("1.5 Vision General", "El usuario puede subir o conectar bases, ejecutar replicas y descargar resultados desde navegador."),
        ("2. Posicionamiento", "Drop Nexus se posiciona como una herramienta academica y tecnica para migracion controlada de datos."),
        ("2.1 Oportunidad de negocio", "Reducir scripts manuales y errores en integraciones entre motores."),
        ("2.2 Definicion del problema", "Las exportaciones manuales no ofrecen trazabilidad ni descarga consistente de la base resultante."),
        ("3. Descripcion de los interesados y usuarios", "Participan usuario tecnico, administrador, docente/evaluador y motores externos."),
        ("3.1 Resumen de los interesados", "Docente, equipo desarrollador y usuarios que requieren replicar datos."),
        ("3.2 Resumen de los usuarios", "Usuarios autenticados que configuran bases y administradores que gestionan usuarios."),
        ("3.3 Entorno de usuario", "Navegador web, API Fastify y servicios Render/Supabase."),
        ("3.4 Perfiles de los interesados", "El docente evalua alcance, el equipo mantiene la app y el usuario ejecuta replicas."),
        ("3.5 Perfiles de los Usuarios", "Usuario tecnico con conocimiento basico de bases y archivos."),
        ("3.6 Necesidades de los interesados y usuarios", "Importar, mapear, replicar, validar errores y descargar resultados."),
        ("4. Vista General del Producto", "Web SaaS con frontend React, backend Fastify y adaptadores por motor."),
    ]:
        doc.add_heading(title, level=1 if title[0].isdigit() and title[1] == "." and title.count(".") == 1 else 2)
        para(doc, text)
    doc.add_heading("4.1 Perspectiva del producto", level=2)
    para(doc, "Producto web integrado con catalogos de archivo y bases externas.")
    doc.add_heading("4.2 Resumen de capacidades", level=2)
    bullets(doc, ["Importar SQL, SQLite, MongoDB, Excel/CSV y .bak.", "Replicar con mapeo.", "Descargar JSON, Excel o SQLite.", "Usar chatbox y skill."])
    doc.add_heading("4.3 Suposiciones y dependencias", level=2)
    para(doc, "Requiere variables de entorno y SQL Server externo para .bak.")
    doc.add_heading("4.4 Costos y precios", level=2)
    para(doc, "Capa gratuita o bajo costo en Render/Supabase.")
    doc.add_heading("4.5 Licenciamiento e instalacion", level=2)
    para(doc, "Instalacion local con npm run install:all dentro de web/.")
    for title, text in [
        ("5. Caracteristicas del producto", "Autenticacion, configuracion, replica por lotes, historial, validaciones y descargas."),
        ("6. Restricciones", "No implementa CDC exact-once; .bak requiere SQL Server externo."),
        ("7. Rangos de calidad", "Seguridad, usabilidad, portabilidad, mantenibilidad y mensajes de error claros."),
        ("8. Precedencia y Prioridad", "Alta prioridad: importar, replicar y descargar. Media: asistencia integrada y extension VS Code."),
        ("9. Otros requerimientos del producto", "Cumplir buenas practicas de seguridad, despliegue en Render y documentacion de uso."),
        ("CONCLUSIONES", "La vision queda centrada en el replicador de datos y no en modulos secundarios."),
        ("RECOMENDACIONES", "Demostrar un flujo completo: importar Excel, replicar y descargar resultado."),
        ("BIBLIOGRAFIA", "Documentacion oficial de Node.js, React, Fastify, Render y Supabase."),
        ("WEBGRAFIA", "Repositorio local web/ y README del proyecto."),
    ]:
        doc.add_heading(title, level=1)
        para(doc, text)


def requirements(doc: Document) -> None:
    sections = [
        ("1. Introduccion", "Este documento especifica los requerimientos de software de Drop Nexus."),
        ("1.1 Proposito", "Definir funcionalidades, restricciones y criterios de aceptacion."),
        ("1.2 Alcance", "Gestion de usuarios, bases, archivos, replicas, historial y descargas."),
        ("1.3 Definiciones, siglas y abreviaturas", "RF: requerimiento funcional. RNF: requerimiento no funcional. Catalogo: base importada normalizada."),
        ("1.4 Referencias", "Documentos FD01, FD02, FD04, README y carpeta web/."),
        ("2. Descripcion general", "Sistema web con frontend, backend y persistencia en Supabase/PostgreSQL."),
        ("2.1 Perspectiva del producto", "Herramienta de replicacion que interactua con motores externos y archivos importados."),
        ("2.2 Funciones del producto", "Importar, configurar, mapear, replicar, monitorear y descargar."),
        ("2.3 Caracteristicas de los usuarios", "Usuarios tecnicos, estudiantes, administradores y evaluadores."),
        ("2.4 Restricciones generales", "Requiere Node.js 20, variables de entorno y SQL Server externo para .bak."),
        ("3. Requerimientos especificos", "Los requerimientos se priorizan por impacto en la replica."),
    ]
    for title, text in sections:
        doc.add_heading(title, level=1 if title.count(".") == 1 else 2)
        para(doc, text)
    doc.add_heading("3.1 Requerimientos funcionales", level=2)
    matrix(doc, ["ID", "Requerimiento", "Prioridad"], [
        ["RF-01", "Autenticar usuarios con JWT.", "Alta"],
        ["RF-02", "Crear configuraciones de bases por conexion o archivo.", "Alta"],
        ["RF-03", "Validar e importar .sql, SQLite, JSON/NDJSON, Excel/CSV y .bak.", "Alta"],
        ["RF-04", "Seleccionar tablas y mapear columnas.", "Alta"],
        ["RF-05", "Ejecutar replicas por lotes con historial.", "Alta"],
        ["RF-06", "Descargar base modificada en JSON, Excel o SQLite.", "Alta"],
        ["RF-07", "Mostrar asistencia contextual mediante chatbox/skill.", "Media"],
    ], [0.7, 4.9, 0.9])
    doc.add_heading("3.2 Requerimientos no funcionales", level=2)
    matrix(doc, ["Categoria", "Requerimiento"], [
        ["Seguridad", "Cifrado AES-256-GCM, JWT, rate limit y aislamiento por usuario."],
        ["Usabilidad", "Mensajes claros para formatos invalidos y .bak no configurado."],
        ["Portabilidad", "Ejecucion local por npm y despliegue en Render."],
        ["Rendimiento", "Replica por lotes y timeout de conexion."],
    ], [1.5, 5.0])
    doc.add_heading("4. Casos de uso", level=1)
    nums(doc, ["Iniciar sesion.", "Importar o configurar base origen.", "Configurar destino.", "Mapear tablas y columnas.", "Ejecutar replica.", "Descargar resultado."])
    doc.add_heading("5. Criterios de aceptacion", level=1)
    bullets(doc, ["Build frontend/backend correcto.", "Excel importable y exportable.", ".bak valida variables faltantes.", "Descarga JSON contiene catalogo actualizado."])
    doc.add_heading("6. Conclusiones", level=1)
    para(doc, "Los requerimientos reflejan la app real implementada en web/.")


def architecture(doc: Document) -> None:
    headings = [
        ("INTRODUCCION", "Se presenta la arquitectura 4+1 simplificada de Drop Nexus."),
        ("Proposito (Diagrama 4+1)", "Describir vistas logica, desarrollo, procesos, despliegue y casos de uso."),
        ("Alcance", "Frontend React, API Fastify, servicios de catalogo/replicacion, adaptadores y despliegue Render."),
        ("Definicion, siglas y abreviaturas", "API, JWT, CDC, Catalogo, Adapter, Job."),
        ("Organizacion del documento", "Objetivos, restricciones, vistas arquitectonicas y conclusiones."),
        ("OBJETIVOS Y RESTRICCIONES ARQUITECTONICAS", "Priorizar seguridad, portabilidad, trazabilidad y facilidad de despliegue."),
        ("Priorizacion de requerimientos", "Alta prioridad para importar, replicar y descargar."),
    ]
    for title, text in headings:
        doc.add_heading(title, level=1 if title.isupper() else 2)
        para(doc, text)
    doc.add_heading("Requerimientos Funcionales", level=2)
    bullets(doc, ["Autenticacion.", "Configuracion de bases.", "Importacion de archivos.", "Replica por lotes.", "Descarga de resultados."])
    doc.add_heading("Requerimientos No Funcionales - Atributos de Calidad", level=2)
    bullets(doc, ["Seguridad mediante cifrado y JWT.", "Portabilidad en Render.", "Mantenibilidad por servicios y adaptadores.", "Usabilidad con validaciones claras."])
    doc.add_heading("Restricciones", level=2)
    para(doc, ".bak depende de SQL Server externo. La replica incremental por offset no reemplaza CDC exact-once.")
    doc.add_heading("REPRESENTACION DE LA ARQUITECTURA DEL SISTEMA", level=1)
    doc.add_heading("Vista de Caso de uso", level=2)
    para(doc, "El usuario inicia sesion, importa bases, configura origen/destino, ejecuta replica y descarga el resultado.")
    doc.add_heading("Diagramas de Casos de uso", level=2)
    para(doc, "Actor Usuario interactua con Gestionar configuracion, Importar archivo, Ejecutar replica y Descargar base modificada.")
    doc.add_heading("Vista Logica", level=2)
    para(doc, "Frontend React consume API Fastify; servicios de dominio coordinan adaptadores de bases y catalogos.")
    doc.add_heading("Diagrama de Subsistemas (paquetes)", level=2)
    matrix(doc, ["Subsistema", "Responsabilidad", "Ubicacion"], [
        ["Frontend", "Interfaz y flujos de usuario", "web/frontend/"],
        ["API", "Rutas y middleware", "web/backend/src/routes/"],
        ["Servicios", "Catalogo, replicacion, usuarios", "web/backend/src/services/"],
        ["Adaptadores", "Conexion a motores", "web/backend/src/adapters/"],
        ["Asistente", "Chatbox, skill y plugin", "web/frontend/src/"],
    ], [1.4, 2.6, 2.5])
    doc.add_heading("Vista de Desarrollo", level=2)
    para(doc, "Codigo TypeScript separado en componentes, paginas, servicios, rutas, adaptadores y utilidades.")
    doc.add_heading("Vista de Procesos", level=2)
    para(doc, "Los jobs de replicacion procesan lotes, guardan progreso y permiten reintentos.")
    doc.add_heading("Vista Fisica o de Despliegue", level=2)
    para(doc, "Render ejecuta backend y sirve frontend/dist; Supabase/PostgreSQL almacena metadatos y catalogos.")
    doc.add_heading("Conclusiones arquitectonicas", level=1)
    para(doc, "La arquitectura es modular, desplegable y alineada con la funcion principal del replicador.")


def final_report(doc: Document) -> None:
    for title, text in [
        ("1. Resumen ejecutivo", "Drop Nexus implementa una web para replicar datos entre bases heterogeneas y descargar la base modificada."),
        ("2. Descripcion del producto final", "La entrega incluye web/, backend, frontend, extension VS Code, skill y documentos actualizados."),
        ("3. Alcance implementado", "Autenticacion, carga de archivos, configuracion, replica, historial, validacion y descarga."),
        ("4. Resultados obtenidos", "La app soporta SQL, SQLite, MongoDB, Excel/CSV y .bak configurado; exporta JSON, Excel y SQLite."),
    ]:
        doc.add_heading(title, level=1)
        para(doc, text)
    doc.add_heading("5. Pruebas realizadas", level=1)
    matrix(doc, ["Prueba", "Resultado"], [
        ["Build backend", "Correcto"],
        ["Build frontend", "Correcto"],
        ["Exportacion JSON", "Correcta"],
        ["Exportacion Excel", "Correcta"],
        ["Importacion Excel", "Correcta"],
        [".bak sin variables", "Error claro"],
    ], [3.0, 3.5])
    for title, text in [
        ("6. Manual breve de ejecucion", "Entrar a web/, ejecutar npm run install:all, configurar variables y levantar backend/frontend."),
        ("7. Limitaciones", "CDC exact-once y restauracion .bak sin SQL Server externo quedan fuera del alcance inmediato."),
        ("8. Conclusiones", "El proyecto cumple con replicar datos y entregar una base modificada descargable."),
        ("9. Recomendaciones", "Presentar un flujo completo con Excel o SQLite y descarga del resultado."),
    ]:
        doc.add_heading(title, level=1)
        para(doc, text)


def proposal(doc: Document) -> None:
    for title, text in [
        ("1. Titulo de la propuesta", PROJECT),
        ("2. Planteamiento del problema", "La transferencia manual de datos entre motores es propensa a errores y carece de trazabilidad."),
        ("3. Justificacion", "La herramienta facilita aprendizaje, pruebas y entrega de resultados descargables."),
        ("4. Objetivo general", "Construir una web para replicar datos entre bases y archivos compatibles."),
        ("5. Objetivos especificos", "Importar fuentes, mapear datos, ejecutar replicas, registrar historial y descargar resultados."),
        ("6. Alcance", "Incluye app web en web/, documentos FD01-FD06, skill, plugin y extension VS Code."),
        ("7. Metodologia", "Desarrollo incremental con validaciones, builds y pruebas funcionales."),
    ]:
        doc.add_heading(title, level=1)
        para(doc, text)
    doc.add_heading("8. Cronograma", level=1)
    matrix(doc, ["Semana", "Actividad", "Entregable"], [
        ["1", "Analisis", "Vision y factibilidad"],
        ["2", "Arquitectura", "Diseño tecnico"],
        ["3-4", "Backend", "API y servicios"],
        ["5", "Frontend", "Interfaz de usuario"],
        ["6", "Asistente", "Skill, plugin y extension"],
        ["7", "Pruebas", "Build y validaciones"],
        ["8", "Documentacion", "Informe final"],
    ], [0.9, 2.8, 2.8])
    doc.add_heading("9. Recursos", level=1)
    para(doc, "Node.js, React, Fastify, Supabase, Render, GitHub y Visual Studio Code.")
    doc.add_heading("10. Criterios de exito", level=1)
    bullets(doc, ["Importar fuente.", "Ejecutar replica.", "Descargar base modificada.", "Documentacion alineada al sistema real."])


def main() -> None:
    build("FD01-EPIS-Informe de Factibilidad.docx", "Informe de Factibilidad", [
        "1. Descripcion del Proyecto", "2. Riesgos", "3. Analisis de la Situacion actual", "4. Estudio de Factibilidad", "5. Analisis Financiero", "6. Conclusiones"
    ], factibilidad)
    build("FD02-EPIS-Informe Vision.docx", "Documento de Vision", [
        "1. Introduccion", "2. Posicionamiento", "3. Descripcion de los interesados y usuarios", "4. Vista General del Producto", "5. Caracteristicas del producto", "6. Restricciones", "7. Rangos de calidad", "8. Precedencia y Prioridad", "9. Otros requerimientos del producto", "CONCLUSIONES", "RECOMENDACIONES", "BIBLIOGRAFIA", "WEBGRAFIA"
    ], vision)
    build("FD03-EPIS-Informe Especificación Requerimientos.docx", "Documento de Especificacion de Requerimientos de Software", [
        "1. Introduccion", "2. Descripcion general", "3. Requerimientos especificos", "4. Casos de uso", "5. Criterios de aceptacion", "6. Conclusiones"
    ], requirements)
    build("FD04-EPIS-Informe Arquitectura de Software.docx", "Documento de Arquitectura de Software", [
        "INTRODUCCION", "OBJETIVOS Y RESTRICCIONES ARQUITECTONICAS", "REPRESENTACION DE LA ARQUITECTURA DEL SISTEMA", "Vista de Caso de uso", "Vista Logica", "Vista de Desarrollo", "Vista de Procesos", "Vista Fisica o de Despliegue"
    ], architecture)
    build("FD05-EPIS-Informe ProyectoFinal.docx", "Informe Final", [
        "1. Resumen ejecutivo", "2. Descripcion del producto final", "3. Alcance implementado", "4. Resultados obtenidos", "5. Pruebas realizadas", "6. Manual breve de ejecucion", "7. Limitaciones", "8. Conclusiones", "9. Recomendaciones"
    ], final_report)
    build("FD06-EPIS-PropuestaProyecto.docx", "Propuesta del Proyecto", [
        "1. Titulo de la propuesta", "2. Planteamiento del problema", "3. Justificacion", "4. Objetivo general", "5. Objetivos especificos", "6. Alcance", "7. Metodologia", "8. Cronograma", "9. Recursos", "10. Criterios de exito"
    ], proposal)


if __name__ == "__main__":
    main()
